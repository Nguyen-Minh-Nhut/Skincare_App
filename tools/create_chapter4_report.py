from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path('/Users/nhut/skincare_app')
OUT = ROOT / 'deliverables'
ASSETS = OUT / 'chapter4_assets'
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

BLUE = '#2457A6'
LIGHT_BLUE = '#EAF2FF'
INK = '#172033'
GRAY = '#5C667A'
GREEN = '#2E7D5B'
ORANGE = '#B76A16'


FONT = '/System/Library/Fonts/Supplemental/Arial.ttf'
FONT_BOLD = '/System/Library/Fonts/Supplemental/Arial Bold.ttf'


def font(size, bold=False):
    path = FONT_BOLD if bold else FONT
    return ImageFont.truetype(path, size)


def canvas(title, size=(2200, 1250)):
    im = Image.new('RGB', size, 'white'); d = ImageDraw.Draw(im)
    d.text((size[0]//2, 45), title, anchor='ma', font=font(42, True), fill=BLUE)
    return im, d


def box(d, xy, text, outline=BLUE, fill='#F6F9FF', fs=27):
    d.rounded_rectangle(xy, radius=28, fill=fill, outline=outline, width=4)
    x1,y1,x2,y2=xy
    d.multiline_text(((x1+x2)//2,(y1+y2)//2), text, anchor='mm', align='center',
                     spacing=8, font=font(fs), fill=INK)


def arrow(d, p1, p2, label=None):
    d.line([p1,p2], fill=GRAY, width=4)
    x1,y1=p1; x2,y2=p2
    import math
    a=math.atan2(y2-y1,x2-x1)
    for da in (2.6,-2.6):
        d.line([(x2,y2),(x2+24*math.cos(a+da),y2+24*math.sin(a+da))],fill=GRAY,width=4)
    if label:
        d.text(((x1+x2)//2,(y1+y2)//2-12),label,anchor='ms',font=font(20),fill=GRAY)


def save_architecture():
    im,d=canvas('KIẾN TRÚC TRIỂN KHAI HIỆN TẠI')
    box(d,(60,400,410,760),'NGƯỜI DÙNG\nChụp / chọn ảnh\nXem và lưu kết quả')
    box(d,(560,300,1050,870),'ỨNG DỤNG FLUTTER\n\nGiao diện\nImage Picker\nTiền xử lý ảnh\nĐiều phối suy luận',fill='#EEF8FF')
    box(d,(1250,220,1600,490),'YOLO11m\nPhát hiện vùng nghi ngờ\nBounding box',outline='#6A55C2',fill='#F3F0FF')
    box(d,(1770,220,2130,490),'MobileNetV2\nXác thực vùng cắt\nMụn / bình thường',outline='#A04F88',fill='#FFF0F8')
    box(d,(1250,690,1600,960),'LƯU TRỮ CỤC BỘ\nSharedPreferences\nTệp ảnh lịch sử',outline=GREEN,fill='#EEFAF4')
    box(d,(1770,690,2130,960),'FIREBASE\nAuthentication\nCloud Firestore',outline=ORANGE,fill='#FFF7EB')
    arrow(d,(410,580),(560,580),'thao tác'); arrow(d,(1050,430),(1250,355),'ảnh'); arrow(d,(1600,355),(1770,355),'vùng cắt')
    arrow(d,(1770,470),(1050,650),'kết quả'); arrow(d,(1050,730),(1250,825),'lưu'); arrow(d,(1050,700),(1770,825),'dữ liệu')
    d.text((1100,1130),'Suy luận trực tiếp trên thiết bị; không truyền ảnh da lên máy chủ riêng.',anchor='mm',font=font(24),fill=GRAY)
    im.save(ASSETS/'architecture.png')


def save_usecase():
    im,d=canvas('SƠ ĐỒ USE CASE CỦA HỆ THỐNG',size=(2000,1400))
    d.rounded_rectangle((350,140,1660,1270),radius=30,fill='#FAFCFF',outline=BLUE,width=4)
    cases=[(520,300,'Đăng nhập / đăng ký'),(1120,300,'Chụp ảnh'),(520,540,'Chọn ảnh từ thiết bị'),(1120,540,'Phân tích tổn thương'),(520,780,'Xem kết quả'),(1120,780,'Lưu / xem / xóa lịch sử'),(820,1040,'Quản lý hồ sơ da')]
    for x,y,t in cases:
        d.ellipse((x-210,y-65,x+210,y+65),fill='#F4F1FF',outline='#6A55C2',width=4); d.text((x,y),t,anchor='mm',font=font(24),fill=INK)
    # actors
    for x,name in [(130,'Người dùng'),(1870,'Quản trị viên')]:
        d.ellipse((x-38,250,x+38,326),outline=INK,width=5); d.line((x,326,x,520),fill=INK,width=5); d.line((x-70,390,x+70,390),fill=INK,width=5); d.line((x,520,x-65,630),fill=INK,width=5); d.line((x,520,x+65,630),fill=INK,width=5); d.text((x,700),name,anchor='mm',font=font(25),fill=INK)
    for y in (300,540,780,1040): arrow(d,(210,430),(350,y))
    arrow(d,(1790,430),(1540,300)); arrow(d,(1790,480),(1540,540))
    im.save(ASSETS/'usecase.png')


def save_sequence():
    im,d=canvas('BIỂU ĐỒ TUẦN TỰ – PHÂN TÍCH ẢNH TRÊN THIẾT BỊ',size=(2200,1400))
    xs=[170,620,1060,1500,1960]; names=['Người dùng','Flutter UI','Tiền xử lý','YOLO11m','MobileNetV2']
    for x,n in zip(xs,names): box(d,(x-130,130,x+130,230),n,fs=24); d.line((x,230,x,1280),fill='#AAB2C0',width=3)
    msgs=[(0,1,340,'1. Chụp/chọn ảnh'),(1,2,460,'2. Đọc bytes, kích thước'),(2,3,580,'3. Ảnh + ngưỡng'),(3,2,700,'4. Bounding box'),(2,4,820,'5. Vùng cắt 224×224'),(4,2,940,'6. Điểm normal/acne'),(2,1,1060,'7. NMS + kết quả'),(1,0,1180,'8. Hiển thị / lưu')]
    for a,b,y,t in msgs: arrow(d,(xs[a],y),(xs[b],y),t)
    im.save(ASSETS/'sequence.png')


def save_data_model():
    im,d=canvas('MÔ HÌNH DỮ LIỆU LOGIC',size=(2200,1350))
    entities=[((50,180,610,530),'users\nuid (PK)\nemail\nusername\nname\nskinType'),((820,180,1380,530),'wardrobe\ndocumentId (PK)\nuserId (FK)\nbarcode\nname\nimage'),((1590,180,2150,530),'notifications\ndocumentId (PK)\nuserId (FK)\ntitle\nisRead\ncreatedAt'),((50,750,610,1100),'posts\ndocumentId (PK)\nuserId (FK)\ncontent\ntag\ncreatedAt'),((820,750,1380,1100),'comments\ndocumentId (PK)\npostId (FK)\nuserId (FK)\ncontent'),((1590,750,2150,1100),'saved_skin_scans (local)\nid (PK)\nuserId\nimagePath\nacneCount\nlevel\ncreatedAt')]
    for xy,t in entities: box(d,xy,t,fill='#F8FAFF',fs=23)
    arrow(d,(610,355),(820,355),'1—N'); arrow(d,(1380,355),(1590,355),'1—N'); arrow(d,(330,530),(330,750),'1—N'); arrow(d,(610,925),(820,925),'1—N'); arrow(d,(610,430),(1590,900),'userId')
    im.save(ASSETS/'data_model.png')


def save_ui_crop():
    source=Path('/var/folders/_y/_9mzv6h54h55zcsm9ggtg53h0000gp/T/TemporaryItems/NSIRD_screencaptureui_UyOGCp/Screenshot 2026-08-22 at 11.56.23.png')
    if source.exists():
        im=Image.open(source)
        # Vùng cửa sổ emulator trong ảnh chụp Android Studio.
        crop=im.crop((1640,145,2320,1650))
        crop.save(ASSETS/'ui_detection.png')


def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)


def set_font(run, size=13, bold=False, color=INK, italic=False):
    run.font.name='Times New Roman'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Times New Roman'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color.replace('#',''))


def add_body(doc, text, bold_prefix=None):
    p=doc.add_paragraph(); p.style='Normal'
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); set_font(r,bold=True)
        r=p.add_run(text[len(bold_prefix):]); set_font(r)
    else: set_font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); set_font(p.add_run(item)); p.paragraph_format.space_after=Pt(3)


def add_numbered(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Number'); set_font(p.add_run(item)); p.paragraph_format.space_after=Pt(4)


def add_caption(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(9)
    set_font(p.add_run(text),size=11,italic=True,color=GRAY)


def add_table(doc, headers, rows, widths_cm):
    table=doc.add_table(rows=1, cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for i,(h,w) in enumerate(zip(headers,widths_cm)):
        c=table.rows[0].cells[i]; c.width=Cm(w); set_cell_shading(c,'DCE8F8'); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(h),size=10.5,bold=True,color=INK)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells=table.add_row().cells
        for i,(value,w) in enumerate(zip(row,widths_cm)):
            cells[i].width=Cm(w); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i in (0,len(row)-1) else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(str(value)),size=10.5,color=INK)
    table.style='Table Grid'
    return table


def add_figure(doc, path, caption, width=6.25):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path),width=Inches(width))
    add_caption(doc,caption)


def build_doc():
    save_architecture(); save_usecase(); save_sequence(); save_data_model(); save_ui_crop()
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(3.0); sec.right_margin=Cm(2.0)
    sec.header_distance=Cm(1.2); sec.footer_distance=Cm(1.2)

    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Times New Roman'; normal.font.size=Pt(13)
    normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; normal.paragraph_format.line_spacing=1.3; normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.first_line_indent=Cm(1.0)
    for name,size,color,before,after in [('Heading 1',16,BLUE,14,8),('Heading 2',14,BLUE,12,6),('Heading 3',13,'1F4D78',8,4)]:
        s=styles[name]; s.font.name='Times New Roman'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color.replace('#','')); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for name in ['List Bullet','List Number']:
        styles[name].font.name='Times New Roman'; styles[name].font.size=Pt(13); styles[name].paragraph_format.left_indent=Cm(.9); styles[name].paragraph_format.first_line_indent=Cm(-.45); styles[name].paragraph_format.line_spacing=1.2

    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_font(header.add_run('CHƯƠNG 4 – TRIỂN KHAI HỆ THỐNG VÀ ỨNG DỤNG'),size=9,color=GRAY)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run('CHƯƠNG 4'),size=18,bold=True,color=BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run('TRIỂN KHAI HỆ THỐNG VÀ ỨNG DỤNG'),size=22,bold=True,color=INK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(14); set_font(p.add_run('Hệ thống hỗ trợ phát hiện và phân tích tình trạng mụn trên ảnh da mặt'),size=14,italic=True,color=GRAY)
    doc.add_paragraph(); doc.add_paragraph()
    for label,value in [('Sinh viên','........................................................'),('Mã số sinh viên','........................................................'),('Giảng viên hướng dẫn','........................................................'),('Thời gian thực hiện','Năm 2026')]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(label+': '),bold=True); set_font(p.add_run(value))
    doc.add_page_break()

    doc.add_heading('4.1. Mô tả chung',level=1)
    add_body(doc,'Ứng dụng được xây dựng nhằm hỗ trợ người dùng quan sát tình trạng mụn trên ảnh da mặt bằng kỹ thuật học sâu. Hệ thống cho phép chụp ảnh bằng camera hoặc chọn ảnh có sẵn, tự động phát hiện các vùng nghi ngờ, xác thực lại từng vùng và hiển thị kết quả trực quan bằng khung đánh dấu. Bên cạnh chức năng phân tích ảnh, ứng dụng còn quản lý hồ sơ da, lịch chăm sóc, tủ sản phẩm, thông báo, cộng đồng tư vấn và lịch sử phân tích.')
    add_body(doc,'Đối tượng sử dụng chính là cá nhân có nhu cầu theo dõi tình trạng da và tham khảo quy trình chăm sóc phù hợp. Quản trị viên được xem là tác nhân phụ, chỉ xuất hiện ở những chức năng quản lý nội dung như mẹo chăm sóc hoặc danh mục sản phẩm. Phạm vi triển khai hiện tại tập trung vào ứng dụng Flutter trên Android; mã nguồn vẫn duy trì khả năng biên dịch đa nền tảng nhưng mô hình AI và thư viện native được kiểm thử chủ yếu trên Android.')
    add_bullets(doc,['Đầu vào: ảnh JPG/PNG được chụp từ camera hoặc chọn từ thư viện thiết bị.','Đầu ra: ảnh có bounding box, số vùng phát hiện, mức đánh giá và gợi ý chăm sóc.','Dữ liệu tài khoản: email, username, hồ sơ da và các dữ liệu cộng đồng được quản lý bằng Firebase.','Lịch sử phân tích: ảnh và metadata kết quả được lưu cục bộ theo từng mã người dùng.'])
    add_body(doc,'Lưu ý y khoa: Kết quả từ ứng dụng chỉ có giá trị hỗ trợ tham khảo và theo dõi cá nhân. Hệ thống không thay thế bác sĩ da liễu, không đưa ra chẩn đoán bệnh và không được dùng làm căn cứ duy nhất để quyết định điều trị.')

    doc.add_heading('4.2. Kiến trúc tổng thể của hệ thống',level=1)
    add_body(doc,'Kiến trúc triển khai hiện tại theo hướng on-device AI. Hai mô hình được đóng gói trong thư mục assets của ứng dụng và suy luận trực tiếp trên thiết bị. Cách tổ chức này giúp giảm độ trễ mạng, hạn chế truyền ảnh da nhạy cảm ra ngoài và cho phép chức năng nhận diện tiếp tục hoạt động khi máy chủ không sẵn sàng.')
    add_figure(doc,ASSETS/'architecture.png','Hình 4.1. Kiến trúc tổng thể của hệ thống. Nguồn: Tác giả đề xuất.')
    add_table(doc,['Thành phần','Vai trò','Công nghệ'],[
        ('Ứng dụng di động','Thu nhận ảnh, điều phối xử lý, hiển thị và lưu kết quả','Flutter/Dart'),
        ('Tiền xử lý','Giải mã ảnh, cắt vùng, đổi kích thước 224×224, chuẩn hóa tensor','package:image'),
        ('YOLO11m','Phát hiện sơ bộ và trả bounding box cùng confidence','Ultralytics YOLO/TFLite'),
        ('MobileNetV2','Xác thực vùng cắt theo hai lớp normal/acne','tflite_flutter'),
        ('Hậu xử lý','Lọc ngưỡng, so sánh điểm và NMS loại khung trùng','Dart'),
        ('Dữ liệu đám mây','Xác thực, hồ sơ, sản phẩm, bài viết và thông báo','Firebase Auth/Firestore'),
        ('Dữ liệu cục bộ','Thiết lập, hội thoại và lịch sử phân tích','SharedPreferences + tệp ảnh'),
    ],[3.2,8.2,4.6])
    add_caption(doc,'Bảng 4.1. Các thành phần chính trong kiến trúc triển khai.')
    add_body(doc,'Trong phương án mở rộng, pipeline AI có thể được đóng gói thành API REST trên FastAPI hoặc Flask. Khi đó ứng dụng gửi multipart/form-data, máy chủ thực hiện suy luận và trả JSON. Tuy nhiên, phương án này chưa phải luồng xử lý của phiên bản mã nguồn hiện tại và chỉ được xem là hướng triển khai khi cần quản lý mô hình tập trung hoặc sử dụng GPU máy chủ.')

    doc.add_heading('4.3. Công nghệ triển khai',level=1)
    add_body(doc,'Các công nghệ huấn luyện mô hình và xử lý ảnh đã được trình bày tại Mục 2.6. Phần này tập trung vào các công nghệ dùng để triển khai ứng dụng và kết nối mô hình với người dùng.')
    add_table(doc,['Công nghệ','Phiên bản/cấu hình','Mục đích'],[
        ('Flutter và Dart','Flutter theo SDK cục bộ; Dart SDK ^3.11.5','Xây dựng giao diện và logic đa nền tảng'),
        ('tflite_flutter','^0.12.1','Nạp và chạy MobileNetV2'),
        ('ultralytics_yolo','^0.6.13','Nạp và suy luận mô hình phát hiện'),
        ('LiteRT native','1.4.0 JNI compatibility','Bổ sung thư viện native cho Android'),
        ('Firebase Auth','^4.15.3','Đăng ký, đăng nhập, đặt lại mật khẩu'),
        ('Cloud Firestore','^4.17.5','Lưu dữ liệu người dùng và nội dung'),
        ('SharedPreferences','^2.5.5','Lưu cấu hình và metadata cục bộ'),
        ('Image Picker','^1.2.2','Chụp hoặc chọn ảnh'),
        ('Provider','^6.1.5+1','Quản lý trạng thái ngôn ngữ và giao diện'),
        ('Git','Quản lý phiên bản','Theo dõi và hợp nhất thay đổi mã nguồn'),
    ],[3.7,4.5,7.8])
    add_caption(doc,'Bảng 4.2. Công nghệ triển khai ứng dụng.')

    doc.add_heading('4.4. Phân tích chức năng và sơ đồ use case',level=1)
    add_figure(doc,ASSETS/'usecase.png','Hình 4.2. Sơ đồ use case của hệ thống. Nguồn: Tác giả đề xuất.',width=5.9)
    add_table(doc,['Mã','Chức năng','Mô tả ngắn'],[
        ('UC01','Đăng nhập/đăng ký','Xác thực bằng email hoặc username, lưu hồ sơ ban đầu.'),
        ('UC02','Chụp ảnh','Mở camera và tiếp nhận ảnh khuôn mặt.'),
        ('UC03','Chọn ảnh','Lấy ảnh có sẵn từ thư viện thiết bị.'),
        ('UC04','Phân tích','Chạy YOLO11m, MobileNetV2 và hậu xử lý NMS.'),
        ('UC05','Xem kết quả','Hiển thị khung phát hiện, số lượng và nhận xét.'),
        ('UC06','Quản lý lịch sử','Lưu, xem chi tiết và xóa kết quả đã lưu.'),
        ('UC07','Quản lý hồ sơ da','Cập nhật loại da, vấn đề da và thông tin cá nhân.'),
        ('UC08','Quản trị nội dung','Thêm/sửa/xóa mẹo và sản phẩm khi có quyền.'),
    ],[1.5,4.0,10.5])
    add_caption(doc,'Bảng 4.3. Mô tả các use case chính.')

    doc.add_heading('4.5. Quy trình xử lý yêu cầu',level=1)
    add_body(doc,'Quy trình phân tích được thực hiện bất đồng bộ để giao diện không bị khóa. Mỗi lần chọn ảnh tạo một mã thế hệ xử lý; nếu người dùng reset hoặc chọn ảnh khác, kết quả của tác vụ cũ sẽ bị bỏ qua. Cơ chế này ngăn kết quả ảnh trước ghi đè lên ảnh đang hiển thị.')
    add_figure(doc,ASSETS/'sequence.png','Hình 4.3. Biểu đồ tuần tự của chức năng phân tích ảnh. Nguồn: Tác giả đề xuất.')
    add_numbered(doc,['Người dùng chọn nguồn ảnh từ camera hoặc thư viện.','Ứng dụng đọc bytes và xác định kích thước ảnh gốc.','YOLO11m phát hiện các vùng nghi ngờ với ngưỡng confidence cấu hình.','Mỗi bounding box được cắt, đổi kích thước 224×224 và chuẩn hóa theo từng kênh màu.','MobileNetV2 trả điểm normal và acne để xác thực detection yếu.','Các detection đạt điều kiện được đưa qua Non-Maximum Suppression với ngưỡng IoU 0,30.','Ứng dụng ánh xạ tọa độ về vùng hiển thị, vẽ bounding box và tạo phần nhận xét.','Khi người dùng chọn lưu, ảnh được sao chép vào thư mục tài liệu ứng dụng và metadata được ghi cục bộ.'])

    doc.add_heading('4.6. Mô tả các chức năng chính',level=1)
    functions=[
        ('4.6.1. Chụp hoặc chọn ảnh','Nhận một ảnh đầu vào từ camera hoặc thư viện.','Nguồn ảnh và quyền truy cập thiết bị.','ImagePicker trả về đường dẫn tạm; ứng dụng đọc bytes và kích thước.','Ảnh hiển thị trong vùng xem trước.','Người dùng hủy, thiếu quyền hoặc tệp không đọc được.'),
        ('4.6.2. Kiểm tra và tiền xử lý ảnh','Chuẩn bị dữ liệu phù hợp cho hai mô hình.','Bytes ảnh JPG/PNG.','Giải mã, kẹp tọa độ bounding box, cắt vùng, resize 224×224 và chuẩn hóa RGB.','Tensor đầu vào dạng [1,3,224,224].','Ảnh hỏng hoặc định dạng không giải mã được.'),
        ('4.6.3. Phát hiện và xác thực tổn thương','Xác định vị trí nghi ngờ và giảm dương tính giả.','Ảnh gốc và ngưỡng confidence.','YOLO phát hiện; MobileNetV2 xác thực vùng yếu; NMS loại khung trùng.','Danh sách AcneBox gồm tọa độ, kích thước và confidence.','Model chưa tải, thư viện native thiếu hoặc suy luận lỗi.'),
        ('4.6.4. Hiển thị kết quả','Trình bày kết quả dễ quan sát.','Danh sách bounding box và kích thước ảnh.','Co giãn tọa độ theo BoxFit, vẽ khung, đếm số vùng và sinh đánh giá.','Ảnh đánh dấu, số nốt và hướng dẫn tham khảo.','Không phát hiện vùng hoặc kích thước hiển thị thay đổi.'),
        ('4.6.5. Lưu và xem lịch sử','Cho phép theo dõi kết quả theo thời gian.','Ảnh, userId, số nốt, mức độ và khuyến nghị.','Sao chép ảnh vào app documents, ghi JSON vào SharedPreferences, lọc theo userId.','Danh sách lịch sử có ảnh và chi tiết kết quả.','Tệp bị xóa, metadata hỏng hoặc thiếu dung lượng.'),
    ]
    for title,purpose,inp,proc,outp,err in functions:
        doc.add_heading(title,level=2)
        for label,val in [('Mục đích: ',purpose),('Dữ liệu đầu vào: ',inp),('Xử lý: ',proc),('Kết quả đầu ra: ',outp),('Trường hợp lỗi: ',err)]: add_body(doc,label+val,bold_prefix=label)

    doc.add_heading('4.7. Thiết kế cơ sở dữ liệu',level=1)
    add_body(doc,'Hệ thống sử dụng mô hình lưu trữ lai. Cloud Firestore lưu dữ liệu cần đồng bộ giữa các thiết bị; SharedPreferences và hệ thống tệp lưu lịch sử phân tích cục bộ. Vì Firestore là cơ sở dữ liệu NoSQL, các “bảng” dưới đây được hiểu là collection/document logic, không áp dụng khóa ngoại cưỡng chế như cơ sở dữ liệu quan hệ.')
    add_figure(doc,ASSETS/'data_model.png','Hình 4.4. Mô hình dữ liệu logic của hệ thống. Nguồn: Tác giả đề xuất.')
    add_table(doc,['Collection/kho dữ liệu','Khóa','Trường quan trọng','Quan hệ'],[
        ('users','uid','email, username, name, skinType, avatarUrl','Gốc của dữ liệu người dùng'),
        ('usernames','username','uid, createdAt','Khóa duy nhất username → user'),
        ('wardrobe','documentId','barcode, name, image','Subcollection của users'),
        ('notifications','documentId','title, body, isRead, createdAt','Subcollection của users'),
        ('posts','documentId','userId, content, tag, likes','Tham chiếu users bằng userId'),
        ('comments','documentId','userId, content, createdAt','Subcollection của posts'),
        ('products','documentId','name, price, category, rating','Danh mục dùng chung'),
        ('saved_skin_scans_v1','id','userId, imagePath, acneCount, level','Danh sách JSON cục bộ'),
    ],[3.5,3.1,6.2,3.2])
    add_caption(doc,'Bảng 4.4. Cấu trúc dữ liệu chính.')
    add_body(doc,'Ảnh lịch sử không được lưu trực tiếp vào Firestore. Ứng dụng sao chép ảnh vào thư mục skin_scans trong vùng dữ liệu riêng của ứng dụng; metadata chỉ lưu đường dẫn nội bộ. Cách làm này giảm chi phí lưu trữ đám mây và hạn chế đưa ảnh da lên mạng, nhưng lịch sử không tự đồng bộ khi đổi thiết bị hoặc cài lại ứng dụng.')

    doc.add_heading('4.8. Thiết kế giao diện ứng dụng',level=1)
    add_body(doc,'Giao diện được thiết kế theo phong cách hiện đại với nền ambient xanh–tím–hồng, các bề mặt kính mờ, màu xanh làm điểm nhấn và thanh điều hướng năm chức năng. Thông tin quan trọng được phân cấp bằng tiêu đề lớn, thẻ tóm tắt và nút hành động rõ ràng. Ứng dụng hỗ trợ chế độ sáng/tối và hai ngôn ngữ Việt–Anh.')
    if (ASSETS / 'ui_detection.png').exists():
        add_figure(doc,ASSETS/'ui_detection.png','Hình 4.5. Giao diện kết quả phân tích và đánh dấu vùng nghi ngờ. Nguồn: Ảnh chụp ứng dụng của tác giả.',width=3.0)
    add_table(doc,['Màn hình','Thành phần chính','Nguyên tắc thiết kế'],[
        ('Đăng nhập/đăng ký','Logo, trường tài khoản, mật khẩu, nút hành động','Tối giản, nhận diện thương hiệu rõ'),
        ('Trang chủ','Nhật ký da, lần soi gần nhất, routine, mẹo','Ưu tiên thông tin theo tần suất sử dụng'),
        ('AI Camera','Khung ảnh, camera/thư viện, trạng thái xử lý','Một hành động chính, phản hồi tiến trình tức thời'),
        ('Kết quả','Bounding box, số nốt, mức độ, khuyến nghị','Kết quả trực quan, cảnh báo không thay thế y khoa'),
        ('Lịch sử','Danh sách ảnh, thời gian, chi tiết, xóa','Dễ so sánh theo thời gian, tách dữ liệu theo user'),
    ],[3.0,7.0,6.0])
    add_caption(doc,'Bảng 4.5. Các màn hình và định hướng giao diện.')

    doc.add_heading('4.9. Triển khai và cấu hình hệ thống',level=1)
    add_body(doc,'Ứng dụng được build dưới dạng APK Android. Cấu hình Android sử dụng Java/Kotlin target 17, bật multidex và đóng gói bổ sung thư viện libtensorflowlite_jni.so để tương thích giữa Ultralytics YOLO, LiteRT và tflite_flutter. Hai tệp trọng số acne_detector.tflite và model_phan_loai_mun.tflite được khai báo trong pubspec.yaml và nạp từ assets khi người dùng truy cập tab AI lần đầu.')
    add_table(doc,['Hạng mục','Cấu hình triển khai'],[
        ('Nền tảng kiểm thử','Android Emulator – Medium Phone API 36.1, kiến trúc arm64'),
        ('Định dạng ảnh','JPG hoặc PNG; chọn ảnh ở chất lượng 100%'),
        ('Đầu vào MobileNetV2','Tensor float32 [1,3,224,224]'),
        ('Đầu ra MobileNetV2','Tối thiểu 2 điểm: normal và acne'),
        ('Ngưỡng YOLO ban đầu','0,08'),
        ('Ngưỡng tin cậy trực tiếp','0,15'),
        ('Ngưỡng loại khung trùng','IoU = 0,30'),
        ('Mất mạng','AI vẫn chạy cục bộ; chức năng Firebase hiển thị lỗi kết nối/timeout'),
        ('Quản lý tài nguyên','Khởi tạo model lười; dispose YOLO và Interpreter khi màn hình bị hủy'),
    ],[5.0,11.0])
    add_caption(doc,'Bảng 4.6. Tham số triển khai chính.')
    add_body(doc,'Đối với Firebase, emulator hoặc thiết bị phải phân giải được identitytoolkit.googleapis.com và firestore.googleapis.com. Ứng dụng đặt timeout cho thao tác đăng nhập để tránh trạng thái chờ vô hạn. Khi Firestore tạm thời lỗi nhưng Firebase Auth vẫn xác nhận người dùng, AuthWrapper giữ người dùng trong ứng dụng thay vì chuyển sai về onboarding.')
    add_body(doc,'Quy trình build đề xuất gồm: chạy dart format, flutter analyze, flutter test, flutter build apk --debug hoặc --release, cài APK lên thiết bị và kiểm tra log runtime. Trước khi phát hành chính thức cần thay applicationId mẫu, cấu hình chữ ký release riêng và không sử dụng debug signing cho bản production.')

    doc.add_heading('4.10. Kiểm thử ứng dụng',level=1)
    add_body(doc,'Kiểm thử được thực hiện ở ba mức: kiểm tra tĩnh bằng Flutter Analyzer, unit/widget test và kiểm thử runtime trên Android emulator. Tại thời điểm lập tài liệu, analyzer không ghi nhận lỗi và bộ test tự động có 13 trường hợp đạt. Bảng dưới đây kết hợp các trường hợp đã xác nhận với các trường hợp cần tiếp tục đo trong đợt nghiệm thu.')
    rows=[
        ('TC01','Chọn ảnh hợp lệ','Ảnh JPG da mặt','Ảnh hiển thị và bắt đầu phân tích','Đã xác nhận','Đạt'),
        ('TC02','Phát hiện ảnh mẫu','anh_da_mun.jpg','Trả bounding box sau lọc trùng','Phát hiện 4 vùng trong lần đo mẫu','Đạt'),
        ('TC03','Lưu lịch sử','Kết quả phân tích hợp lệ','Lưu ảnh và toàn bộ metadata','Đã kiểm tra logic và giao diện','Đạt'),
        ('TC04','Xóa lịch sử','Một SavedScan đã lưu','Xóa metadata và tệp ảnh','Có xử lý đồng bộ hai thành phần','Đạt'),
        ('TC05','Firestore tạm lỗi','Auth hợp lệ, Firestore lỗi','Không chuyển sai về onboarding','Có unit test resolve destination','Đạt'),
        ('TC06','Đổi tab khi có kết quả','Ảnh và kết quả đang hiển thị','Giữ nguyên trạng thái tab','Dùng IndexedStack khởi tạo lười','Đạt'),
        ('TC07','Chọn ảnh mới khi đang chạy','Hai tác vụ liên tiếp','Bỏ kết quả tác vụ cũ','Có generation token','Đạt'),
        ('TC08','Ảnh không hợp lệ','Tệp không giải mã được','Hiển thị thông báo lỗi','Cần kiểm thử nghiệm thu trực tiếp','Chưa đo'),
        ('TC09','Ảnh không có khuôn mặt','Ảnh phong cảnh','Không phát hiện hoặc cảnh báo phù hợp','Cần bổ sung bộ dữ liệu âm','Chưa đo'),
        ('TC10','Mất DNS Firebase','Không phân giải tên miền','Thông báo lỗi kết nối, AI vẫn dùng được','Đã tái hiện trên emulator','Đạt'),
    ]
    add_table(doc,['Mã','Chức năng','Dữ liệu đầu vào','Kết quả mong đợi','Kết quả thực tế','Trạng thái'],rows,[1.2,2.6,3.0,3.5,4.0,1.7])
    add_caption(doc,'Bảng 4.7. Bảng kiểm thử chức năng.')
    doc.add_heading('4.10.1. Chỉ số hiệu năng cần ghi nhận',level=2)
    add_body(doc,'Để báo cáo có thể tái lập, mỗi phép đo nên thực hiện tối thiểu 10 lần sau một lần warm-up. Báo cáo giá trị trung bình, độ lệch chuẩn và giá trị lớn nhất. Không nên ghi số liệu ước lượng khi chưa đo trên thiết bị thật.')
    add_table(doc,['Chỉ số','Cách đo','Đơn vị','Giá trị'],[
        ('Thời gian đọc và giải mã ảnh','Từ lúc chọn ảnh đến khi có ui.Image','ms','Điền sau khi đo'),
        ('Thời gian YOLO11m','Bao quanh lời gọi predict','ms','Điền sau khi đo'),
        ('Thời gian MobileNetV2','Tổng thời gian phân loại các vùng cắt','ms','Điền sau khi đo'),
        ('Thời gian phản hồi tổng thể','Từ thao tác người dùng đến khi hiển thị kết quả','ms','Điền sau khi đo'),
        ('Bộ nhớ cực đại','Android Studio Profiler trong lúc suy luận','MB','Điền sau khi đo'),
        ('Kích thước APK','Dung lượng tệp build đầu ra','MB','Điền sau khi đo'),
    ],[4.2,7.0,1.7,3.1])
    add_caption(doc,'Bảng 4.8. Mẫu ghi nhận hiệu năng.')

    doc.add_heading('4.11. Đánh giá kết quả triển khai',level=1)
    add_body(doc,'Phiên bản hiện tại đã hình thành một quy trình khép kín từ thu nhận ảnh, suy luận hai tầng, hiển thị trực quan đến lưu lịch sử. Điểm mạnh là suy luận trên thiết bị, dữ liệu ảnh không bắt buộc truyền lên máy chủ và giao diện tích hợp nhiều chức năng hỗ trợ chăm sóc da. Những hạn chế còn lại gồm độ chính xác phụ thuộc dữ liệu huấn luyện, kết quả có thể nhạy với ánh sáng/góc chụp, lịch sử chưa đồng bộ đa thiết bị và chưa có bộ kiểm thử ảnh âm đủ lớn.')
    add_body(doc,'Hướng phát triển gồm hiệu chuẩn confidence trên tập kiểm thử độc lập, bổ sung kiểm tra chất lượng ảnh/khuôn mặt, đo hiệu năng trên thiết bị thật, mã hóa hoặc đồng bộ lịch sử theo lựa chọn người dùng, xây dựng API server khi cần quản lý model tập trung, và thực hiện đánh giá với chuyên gia da liễu trước khi mở rộng phạm vi sử dụng.')

    doc.add_heading('Ghi chú sử dụng tài liệu',level=1)
    add_bullets(doc,['Thay các dòng thông tin sinh viên và giảng viên ở trang đầu.','Đối chiếu tên chính xác của kiến trúc mô hình với Chương 3 và hồ sơ huấn luyện.','Bổ sung ảnh giao diện phiên bản cuối nếu giao diện tiếp tục thay đổi.','Đo và điền số liệu hiệu năng trong Bảng 4.8 trước khi nộp.','Đánh số hình/bảng lại nếu ghép chương này vào luận văn hoàn chỉnh.'])

    path=OUT/'Chuong_4_Trien_khai_he_thong_va_ung_dung.docx'
    doc.save(path)
    return path


if __name__=='__main__':
    print(build_doc())
